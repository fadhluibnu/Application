from xml.sax.saxutils import prepare_input_source
import docx
doc = docx.Document('word/052-EMC-2021.docx')  # Creating word reader object.
# data = ""
# fullText = []
# for para in doc.paragraphs:
#     fullText.append(para.text)
#     data = '\n'.join(fullText)
fullText = []
for i in doc.paragraphs:
    fullText.append(i.text)

index = 0
for i in fullText:
    print(index)
    index += 1
